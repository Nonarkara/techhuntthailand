from __future__ import annotations

import json
import re
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib.parse import quote_plus

from docx import Document


COHORT_SPECS = [
    {
        "table_index": 1,
        "id": "mobility-cohort",
        "label": "Mobility Solutions",
        "label_th": "การเดินทางและขนส่ง",
        "label_zh": "交通与移动方案",
        "focus": "Transit, logistics, access, and city-scale movement systems",
        "focus_th": "ระบบการเดินทาง โลจิสติกส์ การเข้าถึง และการเคลื่อนที่ระดับเมือง",
        "focus_zh": "城市级交通、物流、可达性与流动系统",
        "accent": "#cf3a27",
    },
    {
        "table_index": 2,
        "id": "disaster-cohort",
        "label": "Disaster Response",
        "label_th": "การรับมือภัยพิบัติ",
        "label_zh": "灾害应对方案",
        "focus": "Preparedness, sensing, mapping, and emergency coordination",
        "focus_th": "การเตรียมพร้อม การตรวจจับ แผนที่ และการประสานงานฉุกเฉิน",
        "focus_zh": "防灾准备、感知、地图与应急协同",
        "accent": "#1a4fd8",
    },
    {
        "table_index": 3,
        "id": "heritage-cohort",
        "label": "Heritage Tourism",
        "label_th": "การท่องเที่ยวฐานประวัติศาสตร์",
        "label_zh": "历史文化旅游方案",
        "focus": "Culture, local commerce, and destination experience",
        "focus_th": "วัฒนธรรม เศรษฐกิจท้องถิ่น และประสบการณ์ปลายทาง",
        "focus_zh": "文化、本地商业与目的地体验",
        "accent": "#8b5a2b",
    },
    {
        "table_index": 4,
        "id": "health-cohort",
        "label": "Health & Aging",
        "label_th": "สุขภาพและผู้สูงอายุ",
        "label_zh": "健康与老龄化方案",
        "focus": "Care delivery, wellness, and long-term support",
        "focus_th": "การให้บริการสุขภาพ สุขภาวะ และการดูแลระยะยาว",
        "focus_zh": "医疗服务、健康福祉与长期照护",
        "accent": "#1b7a57",
    },
    {
        "table_index": 5,
        "id": "learning-cohort",
        "label": "Learning & Education",
        "label_th": "การเรียนรู้และการศึกษา",
        "label_zh": "学习与教育方案",
        "focus": "Learning access, schools, upskilling, and digital pedagogy",
        "focus_th": "การเข้าถึงการเรียนรู้ โรงเรียน การยกระดับทักษะ และการสอนดิจิทัล",
        "focus_zh": "学习可及性、学校、技能提升与数字教学",
        "accent": "#6c2bd9",
    },
]

SMART_DOMAIN_SPECS = {
    "mobility": {
        "label": "Smart Mobility",
        "label_th": "การเดินทางอัจฉริยะ",
        "label_zh": "智慧交通",
        "icon": "M",
        "accent": "#2157f2",
        "description": "Transportation, logistics, access, and movement infrastructure",
        "description_th": "โครงสร้างพื้นฐานด้านการเดินทาง โลจิสติกส์ การเข้าถึง และการสัญจร",
        "description_zh": "交通、物流、可达性与流动基础设施",
    },
    "environment": {
        "label": "Smart Environment",
        "label_th": "สิ่งแวดล้อมอัจฉริยะ",
        "label_zh": "智慧环境",
        "icon": "E",
        "accent": "#157347",
        "description": "Air, water, waste, land, and sustainability systems",
        "description_th": "ระบบด้านอากาศ น้ำ ขยะ พื้นที่ และความยั่งยืน",
        "description_zh": "空气、水、废弃物、土地与可持续系统",
    },
    "economy": {
        "label": "Smart Economy",
        "label_th": "เศรษฐกิจอัจฉริยะ",
        "label_zh": "智慧经济",
        "icon": "C",
        "accent": "#b7791f",
        "description": "Commerce, finance, marketplaces, and productivity services",
        "description_th": "บริการด้านการค้า การเงิน ตลาดดิจิทัล และผลิตภาพ",
        "description_zh": "商业、金融、交易平台与生产力服务",
    },
    "governance": {
        "label": "Smart Governance",
        "label_th": "การบริหารภาครัฐอัจฉริยะ",
        "label_zh": "智慧治理",
        "icon": "G",
        "accent": "#7c3aed",
        "description": "Public operations, crisis response, and civic coordination",
        "description_th": "การปฏิบัติการภาครัฐ การรับมือวิกฤต และการประสานงานภาคพลเมือง",
        "description_zh": "公共运营、危机应对与市政协同",
    },
    "living": {
        "label": "Smart Living",
        "label_th": "การใช้ชีวิตอัจฉริยะ",
        "label_zh": "智慧生活",
        "icon": "L",
        "accent": "#dd6b20",
        "description": "Health, wellness, home life, and quality-of-life services",
        "description_th": "บริการด้านสุขภาพ สุขภาวะ การอยู่อาศัย และคุณภาพชีวิต",
        "description_zh": "健康、福祉、居家生活与生活质量服务",
    },
    "people": {
        "label": "Smart People",
        "label_th": "คนอัจฉริยะ",
        "label_zh": "智慧人群",
        "icon": "P",
        "accent": "#d53f8c",
        "description": "Education, skills, inclusion, and citizen capability building",
        "description_th": "การศึกษา ทักษะ การมีส่วนร่วม และการพัฒนาศักยภาพประชาชน",
        "description_zh": "教育、技能、包容与公民能力建设",
    },
    "energy": {
        "label": "Smart Energy",
        "label_th": "พลังงานอัจฉริยะ",
        "label_zh": "智慧能源",
        "icon": "N",
        "accent": "#0ea5b7",
        "description": "Power, EV, charging, and energy optimization",
        "description_th": "พลังงาน ไฟฟ้า ยานยนต์ไฟฟ้า การชาร์จ และการเพิ่มประสิทธิภาพพลังงาน",
        "description_zh": "电力、电动车、充电与能源优化",
    },
}

PROGRAM_TIMELINE = [
    {
        "step": "Research Intake",
        "value": "160 entities",
        "detail": "The report consolidates solution providers, agencies, and technology holders from multiple sources.",
    },
    {
        "step": "Incubation",
        "value": "16 startups",
        "detail": "The program opened recruitment and shortlisted startups ready to develop city-scale solutions.",
    },
    {
        "step": "Online Sessions",
        "value": "6 sessions",
        "detail": "Orientation plus five online class sessions were delivered on a dedicated digital platform.",
    },
    {
        "step": "Onsite Workshop",
        "value": "1 day",
        "detail": "A collaborative workshop connected startups and city representatives at True Digital Park.",
    },
    {
        "step": "Business Matching",
        "value": "70-100 people",
        "detail": "Pitching and negotiation zones were arranged for direct conversations in Bangkok.",
    },
]

REFERENCE_PATH = Path(__file__).with_name("reference_atlas.json")

REFERENCE_NAME_ALIASES = {
    "altotechglobal": "altotech",
    "greenery": "greenery",
    "ไทเกอร์โดรน": "thaitigerdrone",
    "vrรับมือภัยพิบัติ": "vrdisastertraining",
    "citydigitaldataplatform": "citydigitaldata",
    "smartemergencypole": "smartemergency",
    "เฮลท์แอทโฮม": "healthathome",
    "caremanagementsquare": "caremanagement",
    "hlab": "hlabinnovative",
    "หุ่นdinsaw": "dinsawrobot",
}

KEYWORD_DOMAINS = [
    (
        "energy",
        [
            "energy",
            "solar",
            "battery",
            "charging",
            "ev",
            "ไฟฟ้า",
            "พลังงาน",
            "แสงอาทิตย์",
        ],
    ),
    (
        "environment",
        [
            "air",
            "water",
            "waste",
            "green",
            "recycle",
            "flood",
            "pollution",
            "forest",
            "carbon",
            "eco",
            "สิ่งแวดล้อม",
            "ขยะ",
            "น้ำ",
            "อากาศ",
            "ไฟป่า",
            "หมอกควัน",
            "ฝุ่น",
            "น้ำท่วม",
        ],
    ),
    (
        "governance",
        [
            "government",
            "civic",
            "emergency",
            "crisis",
            "disaster",
            "security",
            "gis",
            "city digital",
            "city data",
            "citizen",
            "public service",
            "รับมือภัย",
            "ฉุกเฉิน",
            "วิกฤต",
            "เมือง",
            "แผนที่",
            "ความปลอดภัย",
        ],
    ),
    (
        "living",
        [
            "health",
            "doctor",
            "care",
            "hospital",
            "medical",
            "telemedicine",
            "aging",
            "wellness",
            "elder",
            "สุขภาพ",
            "แพทย์",
            "ผู้สูงอายุ",
            "ดูแล",
            "ยา",
            "หมอ",
        ],
    ),
    (
        "people",
        [
            "education",
            "learning",
            "school",
            "course",
            "english",
            "class",
            "teacher",
            "student",
            "skill",
            "การศึกษา",
            "เรียน",
            "คอร์ส",
            "โรงเรียน",
            "นักเรียน",
            "ทักษะ",
        ],
    ),
    (
        "mobility",
        [
            "mobility",
            "transit",
            "transport",
            "bus",
            "car",
            "ride",
            "logistics",
            "shipment",
            "vehicle",
            "traffic",
            "เดินทาง",
            "ขนส่ง",
            "รถ",
            "เดินทาง",
            "โดยสาร",
        ],
    ),
    (
        "economy",
        [
            "market",
            "commerce",
            "accounting",
            "crm",
            "booking",
            "finance",
            "fintech",
            "retail",
            "shop",
            "sale",
            "business",
            "tourism",
            "e-commerce",
            "เศรษฐกิจ",
            "บัญชี",
            "จอง",
            "ขาย",
            "ร้านค้า",
            "ท่องเที่ยว",
            "ธุรกิจ",
        ],
    ),
]

COHORT_FALLBACK_DOMAINS = {
    "mobility-cohort": "mobility",
    "disaster-cohort": "governance",
    "heritage-cohort": "economy",
    "health-cohort": "living",
    "learning-cohort": "people",
}


def normalize_space(value: str) -> str:
    collapsed = re.sub(r"\s+", " ", value.replace("\xa0", " "))
    return collapsed.strip()


def slugify(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    ascii_only = normalized.encode("ascii", "ignore").decode("ascii")
    lowered = ascii_only.lower()
    cleaned = re.sub(r"[^a-z0-9]+", "-", lowered)
    return cleaned.strip("-") or "solution"


def normalize_key(value: str) -> str:
    normalized = unicodedata.normalize("NFKC", value).casefold()
    normalized = normalized.replace("&", "and")
    return re.sub(r"[\W_]+", "", normalized, flags=re.UNICODE)


def discover_source_doc() -> Path:
    candidates = sorted(Path.cwd().glob("*.docx"))
    if not candidates:
        raise FileNotFoundError("No .docx source file found in the working directory.")
    return candidates[0]


def load_reference_atlas() -> tuple[dict[str, dict[str, Any]], dict[str, Any]]:
    if not REFERENCE_PATH.exists():
        return {}, {"solutions": []}

    payload = json.loads(REFERENCE_PATH.read_text(encoding="utf-8"))
    lookup: dict[str, dict[str, Any]] = {}
    for item in payload.get("solutions", []):
        key = normalize_key(item.get("name", ""))
        if key:
            lookup[key] = item
    return lookup, payload


def build_search_url(name: str, organization: str) -> str:
    query = quote_plus(f"{name} {organization} official website")
    return f"https://www.google.com/search?q={query}"


def extract_reference_item(
    reference_lookup: dict[str, dict[str, Any]], name: str
) -> dict[str, Any] | None:
    key = normalize_key(name)
    alias = REFERENCE_NAME_ALIASES.get(key)
    return reference_lookup.get(alias or key)


def infer_domain(
    name: str, description: str, organization: str, cohort_id: str
) -> str:
    searchable = " ".join([name, description, organization]).casefold()

    for domain_id, keywords in KEYWORD_DOMAINS:
        if any(keyword.casefold() in searchable for keyword in keywords):
            return domain_id

    return COHORT_FALLBACK_DOMAINS[cohort_id]


def enrich_solution(
    base_item: dict[str, Any], reference_item: dict[str, Any] | None
) -> dict[str, Any]:
    if reference_item:
        domain_id = reference_item.get("domain") or infer_domain(
            base_item["name"],
            base_item["description"],
            base_item["organization"],
            base_item["cohort"],
        )
        tags = reference_item.get("tags", [])[:3]
        website = reference_item.get("website")
        description_en = reference_item.get("description") or base_item["description_th"]
    else:
        domain_id = infer_domain(
            base_item["name"],
            base_item["description"],
            base_item["organization"],
            base_item["cohort"],
        )
        tags = []
        website = None
        description_en = base_item["description_th"]

    domain = SMART_DOMAIN_SPECS[domain_id]
    search_url = build_search_url(base_item["name"], base_item["organization"])

    return {
        **base_item,
        "description_en": description_en,
        "description_zh": description_en,
        "domain": domain_id,
        "domain_label": domain["label"],
        "domain_label_th": domain["label_th"],
        "domain_label_zh": domain["label_zh"],
        "domain_description": domain["description"],
        "domain_description_th": domain["description_th"],
        "domain_description_zh": domain["description_zh"],
        "domain_icon": domain["icon"],
        "domain_accent": domain["accent"],
        "tags": tags,
        "website": website,
        "search_url": search_url,
        "has_verified_website": bool(website),
    }


def parse_solutions(
    doc: Document, reference_lookup: dict[str, dict[str, Any]]
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]]]:
    solutions: list[dict[str, Any]] = []
    cohorts: list[dict[str, Any]] = []
    domain_counts = {domain_id: 0 for domain_id in SMART_DOMAIN_SPECS}

    for spec in COHORT_SPECS:
        table = doc.tables[spec["table_index"]]
        cohort_items: list[dict[str, Any]] = []
        running_index = 0

        for row in table.rows:
            cells = [normalize_space(cell.text) for cell in row.cells[:4]]
            if len(cells) < 4:
                continue

            raw_order, name, description, organization = cells
            if not name:
                continue

            running_index += 1
            base_item = {
                "id": f"{spec['id']}-{running_index:03d}-{slugify(name)}",
                "order": running_index,
                "source_order": raw_order or str(running_index),
                "name": name,
                "description": description,
                "description_th": description,
                "organization": organization,
                "cohort": spec["id"],
                "cohort_label": spec["label"],
                "cohort_label_th": spec["label_th"],
                "cohort_label_zh": spec["label_zh"],
                "cohort_focus": spec["focus"],
                "cohort_focus_th": spec["focus_th"],
                "cohort_focus_zh": spec["focus_zh"],
                "cohort_accent": spec["accent"],
            }

            enriched = enrich_solution(
                base_item, extract_reference_item(reference_lookup, name)
            )
            solutions.append(enriched)
            cohort_items.append(enriched)
            domain_counts[enriched["domain"]] += 1

        cohorts.append(
            {
                "id": spec["id"],
                "label": spec["label"],
                "label_th": spec["label_th"],
                "label_zh": spec["label_zh"],
                "focus": spec["focus"],
                "focus_th": spec["focus_th"],
                "focus_zh": spec["focus_zh"],
                "accent": spec["accent"],
                "count": len(cohort_items),
            }
        )

    domains = []
    for domain_id, spec in SMART_DOMAIN_SPECS.items():
        domains.append(
            {
                "id": domain_id,
                "label": spec["label"],
                "label_th": spec["label_th"],
                "label_zh": spec["label_zh"],
                "icon": spec["icon"],
                "accent": spec["accent"],
                "description": spec["description"],
                "description_th": spec["description_th"],
                "description_zh": spec["description_zh"],
                "count": domain_counts[domain_id],
            }
        )

    return solutions, cohorts, domains


def parse_mous(doc: Document) -> list[dict[str, str]]:
    table = doc.tables[6]
    items: list[dict[str, str]] = []

    for row in table.rows:
        label = normalize_space(row.cells[0].text)
        if not label:
            continue

        match = re.match(r"^(MOU\d+)\s+(.+?)\s+(\S+)$", label)
        if match:
            code, solution_name, city = match.groups()
        else:
            code, solution_name, city = label, label, ""

        items.append(
            {
                "code": code,
                "solution": solution_name,
                "city": city,
                "label": label,
            }
        )

    return items


def build_payload(source_doc: Path) -> dict[str, Any]:
    doc = Document(source_doc)
    reference_lookup, reference_payload = load_reference_atlas()
    solutions, cohorts, domains = parse_solutions(doc, reference_lookup)
    organizations = sorted({item["organization"] for item in solutions})
    mous = parse_mous(doc)
    verified_websites = sum(1 for item in solutions if item["has_verified_website"])

    return {
        "meta": {
            "title": "Smart City Thailand Tech Hunt Directory",
            "subtitle": "A structured directory built from the research report source file.",
            "source_file": source_doc.name,
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        },
        "stats": {
            "solutions": len(solutions),
            "organizations": len(organizations),
            "cohorts": len(cohorts),
            "domains": len(domains),
            "incubated_startups": 16,
            "online_sessions": 6,
            "onsite_workshops": 1,
            "business_matching_attendance": "70-100",
            "mous": len(mous),
            "verified_websites": verified_websites,
            "search_fallbacks": len(solutions) - verified_websites,
            "reference_matches": sum(
                1 for item in solutions if item["tags"] or item["has_verified_website"]
            ),
            "reference_catalog": len(reference_payload.get("solutions", [])),
        },
        "domains": domains,
        "cohorts": cohorts,
        "timeline": PROGRAM_TIMELINE,
        "mous": mous,
        "solutions": solutions,
    }


def write_browser_data(payload: dict[str, Any], destination: Path) -> None:
    serialized = json.dumps(payload, ensure_ascii=False, indent=2)
    destination.write_text(f"window.DIRECTORY_DATA = {serialized};\n", encoding="utf-8")


def main() -> None:
    source_doc = discover_source_doc()
    payload = build_payload(source_doc)
    write_browser_data(payload, Path("data.js"))
    print(
        "Generated data.js from"
        f" {source_doc.name} with {payload['stats']['solutions']} solutions and"
        f" {payload['stats']['verified_websites']} verified websites."
    )


if __name__ == "__main__":
    main()
